from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "deliverables" / "创AI-开发与应用报告-完整版.docx"

IMG_P2 = ROOT / "assets" / "p2_problem_scene.png"
IMG_P3 = ROOT / "assets" / "p3_learning_loop.png"
IMG_P6 = ROOT / "assets" / "p6_classroom_flow.png"
IMG_SAMPLE = Path("/Users/apple/Documents/GitHub/xedu-client/courses/ai-showcase-exam-2025/01_目标检测与细粒度分类/animal.jpg")


def set_run_font(run, name="仿宋_GB2312", size=12, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold


def add_paragraph(doc, text, *, align=None, font="仿宋_GB2312", size=12, bold=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, name=font, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    set_run_font(r, name="黑体", size={1: 16, 2: 14, 3: 13}.get(level, 12), bold=True)
    return p


def add_figure(doc, img_path: Path, caption: str, width: float = 6.2):
    if not img_path.exists():
        add_paragraph(doc, f"[图片缺失] {caption}", font="仿宋_GB2312", size=11)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, name="仿宋_GB2312", size=10)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "仿宋_GB2312"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    normal.font.size = Pt(12)

    add_paragraph(doc, "创AI案例开发与应用报告完整版", align=WD_ALIGN_PARAGRAPH.CENTER, font="方正小标宋简体", size=16, bold=True)
    add_paragraph(doc, "案例名称：XEdu Client：面向人工智能课堂的一体化实验学习平台", align=WD_ALIGN_PARAGRAPH.CENTER, font="仿宋_GB2312", size=10)
    add_paragraph(doc, "说明：本版为带图完整版。本版将项目主轴统一为四层：实验环境承载、HTML—Blockly—Jupyter 学习主线、XEduHub 完整工具箱、面向教师的课程分享与更新机制。", align=WD_ALIGN_PARAGRAPH.CENTER, font="仿宋_GB2312", size=10)

    add_heading(doc, "一、开发背景", level=1)
    add_paragraph(doc, "在初中人工智能通识课程实施中，本项目主要面向四类核心问题。第一类是实验环境问题。当前人工智能课程往往涉及 Jupyter、Blockly、Python、样例数据、模型资源等多种载体，实验环境搭建复杂、工具切换频繁，教师大量时间耗费在让实验跑起来，学生也难以在同一任务情境中持续完成学习。第二类是学习路径断裂问题。学生常常先接触讲解页或演示材料，再进入图形化练习，最后才尝试代码环境，但这三者之间缺少连续设计，导致“看过了不等于会做、做过了不等于会改、改过了不等于会创新”。第三类是 AI 实验能力分散问题。图像输入、模型调用、流程执行、结果展示等能力分散在不同工具和脚本里，既不利于课堂实施，也不利于教师复用。第四类是课程共享与更新问题。课程资源并不是一次性静态文件，而是会随着教学推进不断修订、补充、再共享的动态内容。教师往往“上一节、改一节、同步一节”，如果缺少统一机制，课程很难积累为可持续运行的平台资源。")
    add_paragraph(doc, "基于上述问题，我们开发了 XEdu Client。它不是单一实验工具，而是一个更像平台的人工智能教学实验环境：一方面承载 Jupyter、Blockly 与本地实验运行环境，降低人工智能实验的部署与使用门槛；另一方面，以 HTML（体验）—Blockly（实践）—Jupyter（应用创新）组织学生的连续学习路径，并结合 XEduHub 这一完整工具箱提供统一的实验能力支撑。同时，项目面向教师建立课程资源的组织、分享、同步更新与持续迭代机制，使课程不再是一次性静态材料，而是可以在教学过程中不断演进的动态资源。")
    add_paragraph(doc, "在开发过程中，项目还引入生成式人工智能辅助需求梳理、交互文案打磨、代码补全、测试样例整理和文档生成，提升了个人与小团队完成跨前后端教育工具开发的效率。对教师开发者而言，这种方式的重要意义在于：即使没有完整的软件工程团队，也可以借助 AI 辅助把教学需求更快转化为可运行的本地工具。正式申报时，可在本节后附上真实使用的国产生成式 AI 工具、提示词片段、对话截图或采纳后的产出对比。")
    add_figure(doc, IMG_P2, "图1 初中人工智能课堂常见问题场景示意图")
    add_paragraph(doc, "图1 主要用于表达“实验环境分散、学习路径断裂、教师需要在多个工具间切换”的问题背景，属于背景问题页辅助图，而非最终核心论文架构图。")

    add_heading(doc, "二、设计与开发", level=1)
    add_heading(doc, "（一）平台/技术选择", level=2)
    add_paragraph(doc, "项目采用“桌面端入口 + 本地服务支撑”的技术路线。前端基于 Electron + Vite 构建统一学习入口，后端采用 Flask + Python 提供课程资源、运行时装配和 Jupyter 管理能力；实践环境使用 JupyterLab/Notebook，可视化学习模块基于 Blockly 实现，人工智能实验流程结合 XEduHub 等教学资源组织。整体方案强调本地可部署、低网络依赖和课堂可控性，适合学校机房和常态课堂使用。技术选择的重点不是栈本身，而是这些技术共同构成了一个能够承载教学、实验、课程资源和更新机制的平台结构。")
    add_paragraph(doc, "从平台能力看，项目可以分为四层。第一层是实验环境承载层，解决 Jupyter、Blockly、Python、样例数据和模型资源分散的问题。第二层是学习路径层，通过 HTML、Blockly、Jupyter 三种载体分别对应体验、实践与应用创新。第三层是工具箱层，依托 XEduHub 把图像输入、模型调用、流程执行和结果呈现等能力纳入统一支撑。第四层是课程与平台机制层，围绕教师备课、课程导入、上传更新、拉取同步和课堂接入建立平台闭环。换句话说，本项目不是简单叠加多种入口，而是在课堂教学视角下把这些能力重新组织成一个完整平台。")
    add_paragraph(doc, "在生成式人工智能赋能开发方面，建议在正式稿中按真实情况补充三类信息：一是使用了哪些国产大模型或 AI 编程工具；二是分别在哪些环节发挥作用，如界面方案讨论、提示文案打磨、代码补全、测试用例整理、申报文稿初稿生成等；三是提供具有代表性的提示词、结果截图和采纳后的修改片段，以体现可复现的开发过程。")
    add_figure(doc, IMG_P3, "图2 XEdu Client 学习闭环结构图")
    add_paragraph(doc, "图2 可作为 HTML—Blockly—Jupyter 学习主线与整体实验闭环的第一版视觉参考，但后续仍建议补充更明确的编号、分区与短标签，升级为论文级主线图。")

    add_heading(doc, "（二）开发过程", level=2)
    add_paragraph(doc, "项目开发大致经历了以下步骤。首先，明确平台目标，把系统定位为“面向课堂实验的教学平台”，而不是单纯的实验启动器。其次，优先解决环境承载问题，让 Jupyter、Blockly、Python 运行与实验资源能够在同一桌面工作台中被统一管理。第三，围绕课堂学习过程设计 HTML—Blockly—Jupyter 主线：HTML 用于任务体验与感知，Blockly 用于流程实践与逻辑理解，Jupyter 用于代码应用与创新拓展。第四，引入 XEduHub 作为完整工具箱，把图像任务、模型调用、流程执行与结果呈现这些 AI 实验能力从分散脚本中抽离出来，形成统一能力底座。第五，补上教师课程机制，支持课程资源的本地创建、云端导入、上传课程、拉取更新和课堂接入，回应“上一节、改一节、同步一节”的真实教学需求。")
    add_paragraph(doc, "在整个迭代过程中，生成式人工智能被用于辅助需求梳理、接口说明生成、测试样例整理和文稿撰写。例如，在需求澄清阶段，可借助 AI 对课堂痛点进行结构化归纳；在开发阶段，可借助 AI 补全局部代码、整理边界情况；在文档阶段，可借助 AI 初步生成描述框架，再由开发者结合真实教学场景进行修订。正式申报时，可在此处插入 [国产AI工具名称]、[典型提示词]、[采纳后的界面或代码片段]、[开发记录截图] 等证据。")

    add_heading(doc, "（三）功能架构", level=2)
    add_paragraph(doc, "项目最终形成了四个相互支撑的功能层。第一，实验环境承载层，围绕课堂实验连续性提供 Jupyter、Blockly、本地运行与资源装配能力。第二，学习路径层，以 HTML（体验）—Blockly（实践）—Jupyter（应用创新）组织学生的连续学习过程。第三，XEduHub 工具箱层，负责图像输入、模型调用、流程执行与结果展示等能力支撑，使 Blockly 与 Jupyter 两侧都能共享统一的 AI 实验基础。第四，课程与平台机制层，提供课程资源管理、云端导入、上传课程、拉取更新、课堂接入以及教师/学生双角色机制，支撑教师侧的课程分享、同步更新与持续迭代。四层共同服务于同一个目标：让平台既能解决实验运行问题，也能承载课程演进与教学协作。")

    add_heading(doc, "三、应用过程与效果", level=1)
    add_paragraph(doc, "该工具适合用于初中人工智能通识课、项目式学习课或校本课程。教师可以先在课程资源中准备实验包，再让学生从讲解页进入 01_目标检测与细粒度分类，理解任务目标和流程逻辑；随后进入 Blockly 工作区观察“检测、裁剪、分类”的串联关系；再切换到 Jupyter，继续阅读和运行对应代码；最后结合结果图或课堂讨论完成总结与复盘。这样，学生不需要在多个无关软件之间频繁切换，而是在统一工作台中完成从认知到实践的闭环。对于教师来说，这种方式也减少了课堂中反复解释“现在打开哪个软件、这个文件在哪里、这一步为什么跑不起来”的时间消耗。更重要的是，课程并不会在一节课结束后静止不变，而是可以在课后继续修订、上传、同步与再使用，使教学资源具备平台化迭代的能力。")
    add_paragraph(doc, "下面这张样例图片来自当前项目内置的目标检测与细粒度分类任务资源，可用于说明课堂任务的输入素材形态。")
    add_figure(doc, IMG_SAMPLE, "图3 目标检测与细粒度分类任务样例输入图片", width=5.4)
    add_paragraph(doc, "从教学效果看，该工具主要带来了三方面改进。第一，降低门槛。学生可先通过积木化方式理解流程，再进入代码实践，减少“看不懂代码就无法开始”的挫败感。第二，增强连续性。课程资源、实验入口和运行结果保留在同一工作台，学生更容易保持任务上下文。第三，提升复现性。教师可以把实验材料打包为结构化课程资源，便于不同班级复用，也便于同一课程在后续教学中持续迭代。")
    add_paragraph(doc, "如果暂时没有完整量化数据，也建议至少补入一段真实课堂应用描述，例如“某班学生先在 Blockly 中完成流程理解，再进入 Notebook 修改参数并观察结果变化”，增强材料的现场感与可信度。待有条件时，还可在此处补充 [应用学校]、[班级数量]、[学生人数]、[实施课时数]、[学生作品示例]、[课堂反馈截图]、[教师访谈摘录] 等真实证据。")
    add_figure(doc, IMG_P6, "图4 课堂应用流程概念图")
    add_paragraph(doc, "图4 用于帮助评审直观理解教师备课、学生进入 Blockly、切换到 Jupyter、结果复盘，以及课后继续更新课程这一完整平台使用路径。")

    add_heading(doc, "四、创新与反思", level=1)
    add_paragraph(doc, "本项目的创新主要体现在四个方面。第一，它不是把多个入口简单并列，而是以 HTML（体验）—Blockly（实践）—Jupyter（应用创新）构造连续学习主线。第二，它不是把 AI 能力分散在零散脚本中，而是结合 XEduHub 形成统一的实验工具箱底座。第三，它不仅关注学生实验能否跑通，也关注教师课程能否被分享、更新、同步和持续迭代，从而把课程资源纳入平台闭环。第四，它在开发过程中引入生成式人工智能，提升教师或小团队跨界开发教育工具的能力。")
    add_paragraph(doc, "同时，项目仍有进一步完善空间。一是不同学校设备条件差异较大，仍需持续优化安装与运维体验；二是课堂应用成效的数据积累还应更加系统；三是若要扩大推广范围，还需继续补齐安装手册、使用手册、课堂案例包和开发记录等材料。未来还可以在现有基础上继续丰富更多可直接导入的课程样例，使不同主题的人工智能课堂都能更快落地。面向教师的课程分享机制也还可以继续加强，例如多教师协作、版本治理和课程模板化复用。")
    add_paragraph(doc, "总体来看，XEdu Client 不是单一课件或演示程序，而是一个以实验环境承载为基础、以 HTML—Blockly—Jupyter 为学习主线、以 XEduHub 为能力工具箱、以教师课程分享与更新机制为平台闭环的人工智能教学实验平台。它既关注学生如何更自然地学习人工智能，也关注教师如何更稳定地组织课堂实施，并让课程资源在教学过程中持续演进，而不只是一次性展示材料。")

    doc.save(OUT)


if __name__ == "__main__":
    main()

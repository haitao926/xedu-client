from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parent
SRC = ROOT / "创AI-开发与应用报告-完整版.md"
OUT = ROOT / "deliverables" / "创AI-开发与应用报告-论文图版.docx"


def set_font(run, *, name: str = "仿宋_GB2312", size: float = 12, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_spacing(p, *, before: float = 0, after: float = 6, line: float = 1.25):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.first_line_indent = Pt(24)


def add_text_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    r = p.add_run(clean_inline(text))
    set_font(r, size=12)
    return p


def add_center_text(doc: Document, text: str, *, size: float = 10, bold: bool = False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(clean_inline(text))
    set_font(r, size=size, bold=bold)
    return p


def clean_inline(text: str) -> str:
    text = text.replace("`", "")
    return text


def add_heading(doc: Document, text: str, level: int):
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(clean_inline(text))
        set_font(r, name="方正小标宋简体", size=18, bold=True)
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(clean_inline(text))
    if level == 2:
        set_font(r, name="黑体", size=15, bold=True)
    elif level == 3:
        set_font(r, name="黑体", size=13, bold=True)
    else:
        set_font(r, name="黑体", size=12, bold=True)


def add_image(doc: Document, image_path: Path):
    if not image_path.exists():
        add_center_text(doc, f"[图片缺失：{image_path}]", size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(5.9))


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(clean_inline(text))
    set_font(r, size=10)


def resolve_image_path(raw: str) -> Path:
    path = Path(raw)
    if path.is_absolute():
        return path
    return (ROOT / raw).resolve()


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)
    section.start_type = WD_SECTION.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = "仿宋_GB2312"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    normal.font.size = Pt(12)
    return doc


def build_docx():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = setup_document()
    lines = SRC.read_text(encoding="utf-8").splitlines()
    pending_image = False

    for raw in lines:
        line = raw.strip()
        if not line:
            continue

        image_match = re.match(r"!\[[^\]]*\]\(([^)]+)\)", line)
        if image_match:
            add_image(doc, resolve_image_path(image_match.group(1)))
            pending_image = True
            continue

        if pending_image and re.match(r"图\d+", line):
            add_caption(doc, line)
            pending_image = False
            continue

        if line.startswith("# "):
            add_heading(doc, line[2:], 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:], 3)
        elif re.match(r"^(作者|单位|申报类别|案例方向)：", line):
            add_center_text(doc, line, size=10)
        elif line.startswith("关键词："):
            p = add_text_paragraph(doc, line)
            p.paragraph_format.first_line_indent = Pt(0)
        else:
            add_text_paragraph(doc, line)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_docx()

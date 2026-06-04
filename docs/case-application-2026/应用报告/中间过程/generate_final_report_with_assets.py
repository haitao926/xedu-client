from __future__ import annotations

import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
CASE_ROOT = ROOT.parent
REPO_ROOT = CASE_ROOT.parents[1]
SRC = ROOT / "中间过程" / "创AI-开发与应用报告-终稿.md"
OUT = ROOT / "创AI-开发与应用报告-终稿.docx"
PPT_SCRIPT = CASE_ROOT / "PPT" / "中间过程" / "build_project_evidence_assets.py"
IMG = ROOT / "图资源"

FIGURES = [
    (
        "基于上述问题，我们开发了 `XEdu Client`。",
        "figure_1_project_background.png",
        "图1 课堂问题与平台回应：用统一工作台承接环境、路径、能力与资源四类需求。",
    ),
    (
        "项目采用“桌面端入口 + 本地服务支撑 + 课程资源结构化管理”的技术路线。",
        "figure_system_framework_clear.png",
        "图2 系统框架：课程资源、平台支撑层与课堂学习链路的关系。",
    ),
    (
        "这样的技术组合并不是简单叠加工具，而是围绕课堂学习过程进行重新编排。",
        "figure_3_learning_path.png",
        "图3 学习主线：从 HTML 讲解体验到 Blockly 流程理解，再到 Jupyter 代码实践。",
    ),
    (
        "在课程资源建设上，项目以文件夹和结构化配置组织课程内容",
        "figure_course_resource_inventory.png",
        "图4 课程资源包结构：讲解页、工作区、Notebook、脚本、素材和输出在同一目录内闭合。",
    ),
    (
        "讲解体验页面模块。该模块以 HTML 页面承载任务情境",
        "figure_course_page_evidence.png",
        "图5 课程讲解页证据：真实课程页面承载任务情境、步骤说明和代码骨架。",
    ),
    (
        "Blockly 可视化理解模块。学生可以通过积木化方式观察人工智能任务流程",
        "figure_blockly_workflow_evidence.png",
        "图6 Blockly 工作流证据：从真实工作区文件提取“读取、检测、提框、分类、展示”的任务链。",
    ),
    (
        "Jupyter 代码实践模块。学生在完成情境体验和流程理解后",
        "figure_notebook_code_evidence.png",
        "图7 Notebook 与 Python 代码证据：同一实验可从 Notebook 入口运行并由 Python 脚本复现。",
    ),
    (
        "在实际应用中，平台已围绕人工智能课堂实验开展试用",
        "figure_5b_collaborative_course.png",
        "图8 课程共建机制：多位教师可围绕同一门课程共享、使用、修订并同步资源。",
    ),
    (
        "以“运动会上的 AI 裁判”等课堂任务为例",
        "figure_sample_output_evidence.png",
        "图9 样例任务运行结果：使用项目内课程包展示输入、检测、裁剪和结果复盘链条。",
    ),
    (
        "生成式人工智能的引入也是本案例的重要特征。",
        "figure_ai_development_process_clean.png",
        "图10 生成式人工智能辅助开发流程：AI 提高整理效率，教师负责判断、取舍和课堂化落地。",
    ),
    (
        "总体来看，`XEdu Client` 不是单一课件或单次演示程序",
        "figure_application_value_clean.png",
        "图11 应用价值归纳：面向学生、教师和课程建设形成连续学习、稳定实施与资源复用价值。",
    ),
]


def set_font(run, *, name: str = "仿宋_GB2312", size: float = 12, bold: bool = False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def clean_inline(text: str) -> str:
    return text.replace("`", "")


def paragraph(doc: Document, text: str, *, first_line=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if first_line:
        p.paragraph_format.first_line_indent = Pt(24)
    r = p.add_run(clean_inline(text))
    set_font(r, size=12)
    return p


def center(doc: Document, text: str, *, size=10, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(clean_inline(text))
    set_font(r, size=size, bold=bold)
    return p


def heading(doc: Document, text: str, level: int):
    if level == 1:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(clean_inline(text))
        set_font(r, name="方正小标宋简体", size=18, bold=True)
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(clean_inline(text))
    set_font(r, name="黑体", size=15 if level == 2 else 13, bold=True)


def add_figure(doc: Document, filename: str, caption: str):
    path = IMG / filename
    if not path.exists():
        center(doc, f"[图片缺失：{filename}]", size=10)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    width = Inches(5.9)
    if filename in {"figure_system_framework_clear.png", "figure_sample_output_evidence.png"}:
        width = Inches(6.2)
    run.add_picture(str(path), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_font(r, size=10)


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.7)
    section.right_margin = Cm(2.7)
    normal = doc.styles["Normal"]
    normal.font.name = "仿宋_GB2312"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    normal.font.size = Pt(12)
    return doc


def build() -> None:
    subprocess.run(["python3", str(PPT_SCRIPT)], check=True, cwd=REPO_ROOT)

    doc = setup_document()
    lines = SRC.read_text(encoding="utf-8").splitlines()
    inserted: set[str] = set()

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            heading(doc, line[2:], 1)
        elif line.startswith("## "):
            heading(doc, line[3:], 2)
        elif line.startswith("### "):
            heading(doc, line[4:], 3)
        elif re.match(r"^(案例名称|作者|单位)：", line):
            center(doc, line, size=10)
        else:
            paragraph(doc, line)

        for anchor, filename, caption in FIGURES:
            if filename not in inserted and anchor in line:
                add_figure(doc, filename, caption)
                inserted.add(filename)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

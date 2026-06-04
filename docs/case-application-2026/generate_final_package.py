from __future__ import annotations

from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parent
TEMPLATE_DIR = Path("/Users/apple/Downloads/通知附件/3.创AI案例征集指南与模板/创AI模板材料")
if not TEMPLATE_DIR.exists():
    TEMPLATE_DIR = Path("/tmp/xedu_case_apply_zip/附件2：案例征集指南与模板/3.创AI案例征集指南与模板/创AI模板材料")
PPT_TEMPLATE = TEMPLATE_DIR / "演示视频PPT模板.pptx"
REPORT_TEMPLATE = TEMPLATE_DIR / "开发与应用报告.doc"
OUT_DIR = ROOT / "deliverables"

NS = {
    "p": "http://schemas.openxmlformats.org/presentationml/2006/main",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
}


def set_run_font(run, name="仿宋_GB2312", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text, *, style=None, align=None, font="仿宋_GB2312", size=12, bold=False, color=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, name=font, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    r = p.add_run(text)
    set_run_font(r, name="黑体", size={1: 16, 2: 14, 3: 13}.get(level, 12), bold=True)
    return p


def remove_paragraphs(tx_body):
    for p in list(tx_body.findall("a:p", NS)):
        tx_body.remove(p)


def build_paragraph_like(template_p, text):
    p = deepcopy(template_p)
    for child in list(p):
        if child.tag in {f'{{{NS["a"]}}}pPr', f'{{{NS["a"]}}}endParaRPr'}:
            continue
        p.remove(child)
    # Remove any runs after the first and keep paragraph properties from the template.
    r_template = template_p.find("a:r", NS)
    if r_template is None:
        r_template = ET.Element(f'{{{NS["a"]}}}r')
    # clear existing runs in the copy
    for r in list(p.findall("a:r", NS)):
        p.remove(r)
    r = deepcopy(r_template)
    for child in list(r):
        if child.tag != f'{{{NS["a"]}}}rPr':
            r.remove(child)
    t = ET.SubElement(r, f'{{{NS["a"]}}}t')
    t.text = text
    p.insert(1 if p.find("a:pPr", NS) is not None else 0, r)
    return p


def set_shape_texts(sp, lines):
    tx_body = sp.find("p:txBody", NS)
    if tx_body is None:
        return
    template_p = tx_body.find("a:p", NS)
    if template_p is None:
        return
    remove_paragraphs(tx_body)
    body_pr = tx_body.find("a:bodyPr", NS)
    lst_style = tx_body.find("a:lstStyle", NS)
    if body_pr is None or lst_style is None:
        return
    # Rebuild paragraphs under the same body/lst style wrapper.
    for line in lines:
        tx_body.append(build_paragraph_like(template_p, line))


def shape_text_map(slide_no):
    return {
        1: {
            3: ["XEdu Client：面向人工智能课堂的", "一体化实验学习平台"],
            4: ["[待填写：学校或单位全称]"],
            5: ["[待填写：作者姓名]"],
            6: ["创AI案例类别：人工智能学习工具"],
        },
        2: {
            1: [
                "面向初中人工智能课堂，聚焦工具切换多、环境配置复杂、学习路径割裂三类问题。",
                "以“目标检测与细粒度分类”任务为主线，统一课程资源、Blockly 和 Jupyter。",
                "帮助学生在同一任务上下文中完成理解、操作、验证与复盘。",
            ],
            3: ["案例概述"],
        },
        3: {
            4: ["课程资源组织", "讲解页、Notebook、Python、图片素材统一管理"],
            5: ["Blockly 可视化理解", "先看懂“检测-裁剪-分类”流程"],
            6: ["Jupyter 代码实践", "把流程理解迁移到代码表达"],
            7: ["本地课堂运行支持", "降低课堂环境依赖，适合机房部署"],
        },
        4: {
            1: ["一、课程资源与课堂接入"],
            4: [
                "1. 教师按课程文件夹组织讲解页、Notebook、Python 文件、Blockly 工作区和实验素材。",
                "2. 学生进入课堂后可直接回到当前课节对应实验，减少找文件和重新配置的时间。",
                "3. 课程资源与课堂接入统一后，课堂时间能更集中地用于讲解与练习。",
            ],
            3: ["实现功能"],
        },
        5: {
            1: ["二、从 Blockly 到 Jupyter 的学习闭环"],
            4: [
                "1. 学生先在 Blockly 中观察任务流程：目标检测、裁剪区域、细粒度分类。",
                "2. 再进入 Jupyter 阅读和运行对应代码，把可视化理解迁移到代码实践。",
                "3. 同一任务上下文贯穿始终，降低“看流程到写代码”的断层感。",
            ],
            3: ["实现功能"],
        },
        6: {
            4: ["适用场景", "初中通识课、项目式学习课、校本课程"],
            5: ["示例任务", "目标检测与细粒度分类、课堂实验连贯演示"],
            6: ["课堂收益", "降低门槛、增强连续性、提高复现性"],
            7: ["推广条件", "本地部署、课程资源包、低网络依赖"],
            3: ["应用情况"],
        },
        7: {
            1: ["一、典型应用过程"],
            4: [
                "1. 教师先在课程资源中准备讲解页、Blockly 工作区、Notebook 和素材图片。",
                "2. 学生从目标案例进入课堂，先看流程，再做积木，再看代码。",
                "3. 最后结合结果图和课堂讨论完成总结与复盘。",
            ],
            3: ["应用情况"],
        },
        8: {
            1: ["二、应用效果、AI 证据与改进"],
            4: [
                "1. 课堂效果：降低人工智能实验起步门槛，增强课堂连续性。",
                "2. 数据占位：[待补：学校/班级/学生人数/课时数]。",
                "3. AI 证据：补入国产 AI 对话截图、提示词片段、采纳对比。",
                "4. 改进方向：继续完善安装手册、使用手册和课堂案例包。",
            ],
            3: ["应用情况"],
        },
        9: {1: ["谢谢！"]},
    }.get(slide_no, {})


def fill_pptx(template_path: Path, out_path: Path):
    with ZipFile(template_path, "r") as zin, ZipFile(out_path, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("ppt/slides/slide") and item.filename.endswith(".xml"):
                slide_no = int(Path(item.filename).stem.replace("slide", ""))
                root = ET.fromstring(data)
                mapping = shape_text_map(slide_no)
                for idx, lines in mapping.items():
                    shapes = root.findall(".//p:sp", NS)
                    if 1 <= idx <= len(shapes):
                        set_shape_texts(shapes[idx - 1], lines)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            zout.writestr(item, data)


def build_info_docx(out_path: Path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "仿宋_GB2312"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    normal.font.size = Pt(12)

    add_paragraph(
        doc,
        "2026年教师人工智能应用案例征集活动\n创AI案例信息表（终稿）",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        font="方正小标宋简体",
        size=16,
        bold=True,
    )
    add_paragraph(doc, "以下内容按创AI / 人工智能学习工具 / 初中 / 课堂申报风格收口。", align=WD_ALIGN_PARAGRAPH.CENTER, font="仿宋_GB2312", size=10)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    fields = [
        ("案例名称", "XEdu Client：面向人工智能课堂的一体化实验学习平台"),
        ("案例类别", "人工智能学习工具"),
        ("作者姓名", "[待填写：作者本人姓名]"),
        ("单位", "[待填写：学校或单位全称]"),
        ("职务/职称", "[待填写：教师/教研员/信息技术教师等]"),
        ("手机", "[待填写]"),
        ("团队成员", "[待填写：如无团队可写“无”]"),
        ("申报学段", "初中"),
        ("解决的教学问题", "初中人工智能课程中，Notebook 编程、可视化积木和课程资源常分散在不同工具里，环境配置复杂、切换频繁，学生难以在同一任务上下文中持续完成理解、操作、验证与复盘。"),
        ("开发平台/工具", "Electron、Vite、Flask、Python、JupyterLab/Notebook、Blockly、XEduHub、HTML/CSS/JavaScript，以及[待补：实际使用的国产生成式AI工具名称]。"),
        ("特色与创新", "面向初中人工智能课堂，把 Blockly 可视化理解与 Jupyter 代码实践整合到同一工作台，支持课程资源、课堂接入和实验联动，并以本地部署降低课堂环境依赖。"),
        ("相关网址", "代码仓库：[待填写]\n演示地址/视频地址：[待填写；如仅本地运行可写“本地部署演示”]"),
        ("配套资源", "完整代码、应用文档、安装手册、课堂案例包、视频脚本、开发记录"),
        ("案例内容简介", "XEdu Client 是一款面向初中人工智能课堂实验的学习工具，围绕学生“继续完成实验”这一真实场景，将 Jupyter 代码实践、Blockly 可视化编程、课程资源管理、课堂接入与 AI 辅助能力整合到统一工作台。项目重点解决环境搭建难、工具切换多、学习路径割裂、课堂上下文难保留等问题。学生可先通过 Blockly 低门槛理解“目标检测与细粒度分类”等任务流程，再进入 Jupyter 继续代码实验；教师可管理课程资源、预演实验流程并组织课堂实施。项目支持本地部署，适合学校机房与常态课堂使用。"),
    ]
    for left, right in fields:
        row = table.add_row().cells
        row[0].text = left
        row[1].text = right
        for cell in row:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_run_font(r, name="仿宋_GB2312", size=11)
        row[0].paragraphs[0].runs[0].bold = True

    add_paragraph(doc, "作者声明：该案例为本人原创，不涉及抄袭或侵犯他人著作权等问题。", font="仿宋_GB2312", size=11)
    add_paragraph(doc, "作者签名：____________________    年  月  日", font="仿宋_GB2312", size=11)
    add_paragraph(doc, "单位意见：同意上报    单位（盖章）：____________________    年  月  日", font="仿宋_GB2312", size=11)

    doc.save(out_path)


def build_report_docx(out_path: Path):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "仿宋_GB2312"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    normal.font.size = Pt(12)

    add_paragraph(doc, "创AI案例开发与应用报告终稿", align=WD_ALIGN_PARAGRAPH.CENTER, font="方正小标宋简体", size=16, bold=True)
    add_paragraph(doc, "案例名称：XEdu Client：面向人工智能课堂的一体化实验学习平台", align=WD_ALIGN_PARAGRAPH.CENTER, font="仿宋_GB2312", size=10)
    add_paragraph(doc, "说明：以下文本可直接填入官方《开发与应用报告》模板。", align=WD_ALIGN_PARAGRAPH.CENTER, font="仿宋_GB2312", size=10)

    sections = [
        ("一、开发背景", [
            "在初中人工智能通识课程实施中，教师常常面临三个问题：一是讲解材料、可视化练习和代码实践分散在不同工具中；二是环境配置复杂，课堂时间容易耗费在切换和排错上；三是学生在不同入口之间频繁跳转，任务上下文难以连续保留。实际教学中，教师往往需要先用网页或课件讲解任务，再切换到图形化工具帮助学生理解流程，最后再回到 Notebook 或 Python 环境做代码验证。这样不仅打断课堂节奏，也容易让基础较弱的学生在“看懂任务”和“进入实践”之间失去衔接。",
            "针对这些问题，我们开发了 XEdu Client。该工具把 Jupyter 代码实验、Blockly 可视化理解、课程资源管理和课堂接入整合到同一学习工作台，帮助学生在一个连续任务中完成“理解原理、观察流程、进入代码、验证结果、回顾总结”的学习闭环。以 01_目标检测与细粒度分类 为例，学生可先理解“先检测、再裁剪、再分类”的任务逻辑，再通过 Blockly 观察流程编排，最后进入 Jupyter 阅读和运行代码，从而把抽象的人工智能流程转化为可操作、可验证的课堂实验活动。",
            "在开发过程中，项目还引入生成式人工智能辅助需求梳理、交互文案打磨、代码补全、测试样例整理和文档生成，提升了个人与小团队完成跨前后端教育工具开发的效率。对教师开发者而言，这种方式的重要意义在于：即使没有完整的软件工程团队，也可以借助 AI 辅助把教学需求更快转化为可运行的本地工具。正式申报时，请在本节补入真实使用的国产生成式 AI 工具、提示词片段、对话截图或采纳后的产出对比。",
        ]),
        ("二、设计与开发", []),
        ("（一）平台/技术选择", [
            "项目采用“桌面端入口 + 本地服务支撑”的技术路线。前端基于 Electron + Vite 构建统一学习入口，后端采用 Flask + Python 提供课程资源、运行时装配和 Jupyter 管理能力；实践环境使用 JupyterLab/Notebook，可视化学习模块基于 Blockly 实现，人工智能实验流程结合 XEduHub 等教学资源组织。整体方案强调本地可部署、低网络依赖和课堂可控性，适合学校机房和常态课堂使用。",
            "从教学适配性看，这一组合兼顾了初中学生的学习起点和人工智能课程的实践需求。Blockly 适合理解流程，Jupyter 适合支撑代码阅读、运行和修改，课程资源系统则把讲解页、样例图片、积木工作区和 Notebook 入口组织在同一课程结构中，减少认知负担。这样的组合并不是简单叠加多种工具，而是针对课堂教学把不同学习阶段的载体重新编排：在理解阶段强调“看得懂”，在实践阶段强调“做得成”，在复盘阶段强调“回得去”。",
            "在生成式人工智能赋能开发方面，建议在正式稿中按真实情况补充三类信息：一是使用了哪些国产大模型或 AI 编程工具；二是分别在哪些环节发挥作用，如界面方案讨论、提示文案打磨、代码补全、测试用例整理、申报文稿初稿生成等；三是提供具有代表性的提示词、结果截图和采纳后的修改片段，以体现可复现的开发过程。",
        ]),
        ("（二）开发过程", [
            "项目开发大致经历了以下步骤。首先，明确产品主线，把系统定位为面向课堂实验的一体化学习工具，而不是单纯的实验启动器。其次，梳理课堂任务链路，确定“讲解页—Blockly—Jupyter—结果展示”这一基本学习闭环，并以 01_目标检测与细粒度分类 作为代表样例反复验证路径是否顺畅。再次，打通 Jupyter 与 Blockly 双入口学习路径，让学生既能通过可视化积木理解人工智能任务流程，也能继续进入 Notebook/Python 做代码实践。随后，组织课程资源与课堂接入，支持教师以课程文件夹方式管理讲解页、Notebook、Python 文件、Blockly 工作区、图片和数据等实验材料，并在课堂接入后让学生直接回到当前课节对应实验。之后，构建 Blockly 可运行学习模块，使其能够承载真实课堂任务，包括流程展示、结果呈现和与代码入口的联动。最后，接入 Jupyter 管理能力，将部分复杂度从学生侧转移到工具侧，提高课堂使用的稳定性与一致性。",
            "在整个迭代过程中，生成式人工智能被用于辅助需求梳理、接口说明生成、测试样例整理和文稿撰写。例如，在需求澄清阶段，可借助 AI 对课堂痛点进行结构化归纳；在开发阶段，可借助 AI 补全局部代码、整理边界情况；在文档阶段，可借助 AI 初步生成描述框架，再由开发者结合真实教学场景进行修订。正式申报时，可在此处插入 [国产AI工具名称]、[典型提示词]、[采纳后的界面或代码片段]、[开发记录截图] 等证据，突出“AI 赋能开发”的特征。",
        ]),
        ("（三）功能架构", [
            "项目最终形成了四个相互支撑的功能模块。第一，学生实验工作台，围绕课堂实验连续性提供 Jupyter 与 Blockly 双入口。第二，课程资源管理模块，教师可按课程目录组织讲解页、Notebook、Python、Blockly、图片和说明文档，形成结构化实验包。第三，Blockly 可运行学习模块，学生可通过拖拽积木理解 AI 任务流程、查看运行结果，并逐步过渡到代码表达。第四，本地运行与教学支持模块，系统负责 Jupyter 运行、资源装配、接口服务及教师/学生模式切换等底层支撑，降低课堂使用门槛。四个模块共同服务于同一个目标：让学生在同一任务情境下持续完成学习，而不是在多个无关工具之间来回迁移。",
        ]),
        ("三、应用过程与效果", [
            "该工具适合用于初中人工智能通识课、项目式学习课或校本课程。教师可以先在课程资源中准备实验包，再让学生从讲解页进入 01_目标检测与细粒度分类，理解任务目标和流程逻辑；随后进入 Blockly 工作区观察“检测、裁剪、分类”的串联关系；再切换到 Jupyter，继续阅读和运行对应代码；最后结合结果图或课堂讨论完成总结与复盘。这样，学生不需要在多个无关软件之间频繁切换，而是在统一工作台中完成从认知到实践的闭环。对于教师来说，这种方式也减少了课堂中反复解释“现在打开哪个软件、这个文件在哪里、这一步为什么跑不起来”的时间消耗。",
            "从教学效果看，该工具主要带来了三方面改进。第一，降低门槛。学生可先通过积木化方式理解流程，再进入代码实践，减少“看不懂代码就无法开始”的挫败感。第二，增强连续性。课程资源、实验入口和运行结果保留在同一工作台，学生更容易保持任务上下文。第三，提升复现性。教师可以把实验材料打包为结构化课程资源，便于不同班级复用，也便于同一课程在后续教学中持续迭代。",
            "正式申报时，建议在本节补充真实应用数据，如 [应用学校]、[班级数量]、[学生人数]、[实施课时数]、[学生作品示例]、[课堂反馈截图]、[教师访谈摘录]。如果有对比数据，也可进一步说明该工具在课堂组织效率、学生参与度、实验完成率等方面的提升情况。若暂时没有完整量化数据，也建议至少补入一段真实课堂应用描述，例如“某班学生先在 Blockly 中完成流程理解，再进入 Notebook 修改参数并观察结果变化”，增强材料的现场感与可信度。",
        ]),
        ("四、创新与反思", [
            "本项目的创新主要体现在四个方面。第一，围绕“课堂实验继续完成”重组完整学习链条，不再把人工智能课堂拆成彼此割裂的讲解、演示和编程环节。第二，把 Blockly 与 Jupyter 放在同一教学工作台中，让可视化理解与代码实践自然衔接。第三，强调本地部署与资源打包，适配学校机房和普通教室环境。第四，在开发过程中引入生成式人工智能，提升教师或小团队跨界开发教育工具的能力。",
            "同时，项目仍有进一步完善空间。一是不同学校设备条件差异较大，仍需持续优化安装与运维体验；二是课堂应用成效的数据积累还应更加系统；三是若要扩大推广范围，还需继续补齐安装手册、使用手册、课堂案例包和开发记录等材料。未来还可以在现有基础上继续丰富更多可直接导入的课程样例，使不同主题的人工智能课堂都能更快落地。",
            "总体来看，XEdu Client 不是单一课件或演示程序，而是一款面向真实教学问题开发的人工智能学习工具。它既关注学生如何更自然地学习人工智能，也关注教师如何更稳定地组织课堂实施。正因为它兼顾了“课堂可用”和“学习可持续”两个维度，所以更接近一项可推广、可复用、可继续演进的教育工具成果，而不只是一次性的展示材料。",
        ]),
    ]

    for title, paras in sections:
        add_heading(doc, title, level=1)
        for para in paras:
            add_paragraph(doc, para, font="仿宋_GB2312", size=12)

    doc.save(out_path)


def build_manifest(out_dir: Path):
    manifest = out_dir / "README.md"
    manifest.write_text(
        "# 创AI 申报终稿包\n\n"
        "- `创AI-案例信息表-终稿.docx`\n"
        "- `创AI-开发与应用报告-终稿.docx`\n"
        "- `创AI-演示视频PPT-终稿.pptx`\n\n"
        "说明：PPT 已按官方 9 页模板填充文本，Word 报告已按官方章节结构生成。\n",
        encoding="utf-8",
    )


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_info_docx(OUT_DIR / "创AI-案例信息表-终稿.docx")
    build_report_docx(OUT_DIR / "创AI-开发与应用报告-终稿.docx")
    fill_pptx(PPT_TEMPLATE, OUT_DIR / "创AI-演示视频PPT-终稿.pptx")
    build_manifest(OUT_DIR)


if __name__ == "__main__":
    main()
